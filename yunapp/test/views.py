# -*- coding: utf-8 -*-

from flask import Blueprint, render_template, send_file, jsonify
import os, io, StringIO
from flask import Flask, request, redirect, url_for
<<<<<<< HEAD
from yunapp.yunapps import app
from docx import Document
from docx.shared import Inches
from HTMLParser import HTMLParser
=======
from yunapp.yunapps import app, yun_redis
>>>>>>> ac2fb8168ddcce8a084c43b567cb98ac72e39358

ALLOWED_EXTENSIONS = set(['txt', 'pdf', 'png', 'jpg', 'jpeg', 'gif'])

test = Blueprint('test', __name__)

UPLOAD_FOLDER = '/home/seanwu/uploads/'
# app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# create a subclass and override the handler methods
class MyHTMLParser(HTMLParser):
    p = None

    def __init__(self, doc_file):
        self.doc_file = doc_file
        self.pre_end_tag=''
        self.document = Document()
        HTMLParser.__init__(self)
        self.pre_tag = ''
        self.current_tag = ''

    def handle_starttag(self, tag, attrs):
        self.pre_tag = self.current_tag
        self.current_tag = tag

    def handle_endtag(self, tag):
        self.pre_end_tag=self.current_tag
        self.current_tag = self.pre_tag

    def handle_data(self, data):
        if 'h1' in self.current_tag:
            self.document.add_heading(data, level=1)
        elif 'h2' in self.current_tag:
            self.document.add_heading(data, level=2)
        elif 'h3' in self.current_tag:
            self.document.add_heading(data, level=3)
        elif 'p' == self.current_tag:
            self.pre_end_tag
            if self.pre_end_tag == 'span':
                self.p.add_run(data)
            else:
                self.p = self.document.add_paragraph(data)
        elif 'span' == self.current_tag:
            # self.p.add_run(data).underline=True
            self.p.add_run(data).bold = True
        elif 'li' in self.current_tag:
            self.document.add_paragraph(data, style='ListBullet')
        else:
            print 'lalala chucuola~~~'

    def get_file(self):
        self.document.save(self.doc_file)
        return self.doc_file.getvalue()


@test.route('/')
def hello_world():
    return render_template('index.html')


@test.route('/doctest/<filename>.docx')
def get_doc(filename):
    doc_file = StringIO.StringIO()
    parser = MyHTMLParser(doc_file)
    parser.feed(
        '<h1>WYSIHTML5 - A better approach to rich text editing</h1><p>wysihtml5 is an <span>open source</span> rich text editor based on HTML5 technology and the progressive-enhancement approach. It uses a sophisticated security concept and aims to generate fully valid HTML5 markup by preventing unmaintainable tag soups and inline styles.</p><h2>Features</h2><p>alsfkjalskfj;alk <span>open source</span>sfj;alksfja;slkfdj</p><li>biaotibiaotibiaoti</li>')

    # return 'hello world'
    # return send_file(io.BytesIO(fileData))
    s = io.BytesIO(parser.get_file())
    doc_file.close()
    parser.close()
    return send_file(s)


@test.route('/file/upload', methods=['POST'])
def file_upload():
    if request.method == 'POST':
        file = request.files['userfile']
        if file and allowed_file(file.filename):
            filename = file.filename
            file.save(os.path.join(UPLOAD_FOLDER, filename))
            return {"fid": "000", "fname": '+filename+'}
        else:
            return {"fid": "000", "errMsg": 'error'}


@test.route('/<path:filename>')
def template_load(filename=None):
    if not filename:
        return render_template('index.html')
    else:
        return render_template('/test/' + filename)


def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1] in ALLOWED_EXTENSIONS

@test.route('/add_user_visit', methods=['GET'])
def add_user_redis():
    yun_redis.incr('user:visit', 1)
    return jsonify({'success': True})

@test.route('/get_user_visit', methods=['GET'])
def get_user_visit():
    print yun_redis.RESPONSE_CALLBACKS
    return jsonify({'user_visit': yun_redis.get('user:visit')})
