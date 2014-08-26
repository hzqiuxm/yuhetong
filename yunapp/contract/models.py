# -*- coding: utf-8 -*-
__author__ = 'Seanwu'
from HTMLParser import HTMLParser
from docx import Document

class Yunhetong_HTMLParser(HTMLParser):
    p = None

    def __init__(self, doc_file):
        self.doc_file = doc_file
        self.pre_end_tag=''
        self.document = Document()
        HTMLParser.__init__(self)
        self.pre_tag = ''
        self.current_tag = ''

    def handle_starttag(self, tag, attrs):
        self.pre_tag = self.current_tag
        self.current_tag = tag

    def handle_endtag(self, tag):
        self.pre_end_tag=self.current_tag
        self.current_tag = self.pre_tag

    def handle_data(self, data):
        if 'h1' in self.current_tag:
            self.document.add_heading(data, level=1)
        elif 'h2' in self.current_tag:
            self.document.add_heading(data, level=2)
        elif 'h3' in self.current_tag:
            self.document.add_heading(data, level=3)
        elif 'p' == self.current_tag:
            self.pre_end_tag
            if self.pre_end_tag == 'span':
                self.p.add_run(data)
            else:
                self.p = self.document.add_paragraph(data)
        elif 'span' == self.current_tag:
            # self.p.add_run(data).underline=True
            self.p.add_run(data).bold = True
        elif 'li' in self.current_tag:
            self.document.add_paragraph(data, style='ListBullet')
        else:
            print 'lalala chucuola~~~'

    def get_file(self):
        self.document.save(self.doc_file)
        return self.doc_file.getvalue()
